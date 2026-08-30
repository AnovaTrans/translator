"""Universal document <-> XLIFF conversion.

Any non-XLIFF format is turned INTO an XLIFF (source-filled, empty targets)
so the one translation engine handles everything, then the translations are
written BACK into a copy of the original document — the platform's canonical
"everything is XLIFF internally" flow.

v1 formats: .txt, .docx. PPTX/XLSX/HTML follow the same extract/rebuild shape.
DOCX formatting is best-effort (paragraph-level; the whole translation goes in
the paragraph's first run, keeping its style). High-fidelity inline formatting
is the platform's Document-Preparation (OCR->XLIFF) step's job.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import Dict, List, Tuple
from xml.sax.saxutils import escape

SUPPORTED_DOC_EXT = {".txt", ".docx"}


def ext_of(filename: str) -> str:
    return Path(filename).suffix.lower()


def is_document(filename: str) -> bool:
    return ext_of(filename) in SUPPORTED_DOC_EXT


# ---------------- XLIFF envelope ----------------

def _build_xliff(pairs: List[Tuple[str, str]], src: str, tgt: str, original: str = "document") -> bytes:
    units = "".join(
        f'<trans-unit id="{escape(str(i))}"><source>{escape(s)}</source><target></target></trans-unit>'
        for i, s in pairs
    )
    return (
        f'<?xml version="1.0" encoding="UTF-8"?>\n'
        f'<xliff version="1.2">\n'
        f'<file source-language="{escape(src)}" target-language="{escape(tgt)}" '
        f'datatype="plaintext" original="{escape(original)}">\n<body>\n{units}\n</body>\n</file>\n</xliff>'
    ).encode("utf-8")


# ---------------- TXT ----------------

def _decode(b: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "utf-16", "latin-1", "cp1252"):
        try:
            return b.decode(enc)
        except Exception:
            continue
    return b.decode("utf-8", errors="ignore")


def _txt_pairs(text: str) -> List[Tuple[str, str]]:
    return [(str(i), ln.strip()) for i, ln in enumerate(text.splitlines()) if ln.strip()]


def _txt_rebuild(text: str, translations: Dict[str, str]) -> bytes:
    out = []
    for i, ln in enumerate(text.splitlines()):
        out.append(translations.get(str(i), ln) if ln.strip() else ln)
    return ("\n".join(out)).encode("utf-8")


# ---------------- DOCX ----------------

def _docx_walk(document):
    """Paragraphs in a stable order: body paragraphs, then every table cell's
    paragraphs. Extract and rebuild use the same walk, so ids line up."""
    paras = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    return paras


def _set_para_text(p, text: str) -> None:
    """Replace a paragraph's text, keeping its first run's formatting/style."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _docx_pairs(file_bytes: bytes) -> List[Tuple[str, str]]:
    import docx
    document = docx.Document(io.BytesIO(file_bytes))
    return [(str(i), p.text) for i, p in enumerate(_docx_walk(document)) if p.text.strip()]


def _docx_rebuild(file_bytes: bytes, translations: Dict[str, str]) -> bytes:
    import docx
    document = docx.Document(io.BytesIO(file_bytes))
    for i, p in enumerate(_docx_walk(document)):
        if p.text.strip() and str(i) in translations:
            _set_para_text(p, translations[str(i)])
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------------- public API ----------------

def to_xliff(file_bytes: bytes, filename: str, src_lang: str, tgt_lang: str) -> bytes:
    """Convert a supported document to a source-filled XLIFF."""
    ext = ext_of(filename)
    if ext == ".txt":
        pairs = _txt_pairs(_decode(file_bytes))
    elif ext == ".docx":
        pairs = _docx_pairs(file_bytes)
    else:
        raise ValueError(f"Unsupported document format: {ext}")
    return _build_xliff(pairs, src_lang, tgt_lang, original=filename)


def parse_reference_chunks(content: bytes, filename: str) -> List[str]:
    """Parse a target-language reference document into text chunks for the
    semantic style matcher. TXT / DOCX / PDF; ignores very short lines."""
    ext = ext_of(filename)
    chunks: List[str] = []
    try:
        if ext == ".txt":
            for ln in _decode(content).replace("\r\n", "\n").split("\n"):
                ln = ln.strip()
                if len(ln) > 15:
                    chunks.append(ln)
        elif ext == ".docx":
            import docx
            for p in docx.Document(io.BytesIO(content)).paragraphs:
                if len(p.text.strip()) > 15:
                    chunks.append(p.text.strip())
        elif ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    for ln in (page.extract_text() or "").split("\n"):
                        if len(ln.strip()) > 15:
                            chunks.append(ln.strip())
    except Exception:
        pass
    return chunks


def rebuild(translations: Dict[str, str], original_bytes: bytes, filename: str) -> Tuple[bytes, str]:
    """Write translations back into a copy of the original document. Returns
    (bytes, output_filename)."""
    ext = ext_of(filename)
    stem = Path(filename).stem
    if ext == ".txt":
        return _txt_rebuild(_decode(original_bytes), translations), f"{stem}_translated.txt"
    if ext == ".docx":
        return _docx_rebuild(original_bytes, translations), f"{stem}_translated.docx"
    raise ValueError(f"Unsupported document format: {ext}")
